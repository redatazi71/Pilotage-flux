"""Génère docs/cadrage_v4.docx à partir de docs/cadrage_v4.md.

Convertit le Markdown en DOCX en respectant la mise en page du cadrage
v3 original (Heading 1 / Heading 2 / Heading 3 / Normal + tableaux).

Usage : python docs/build_cadrage_v4_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "cadrage_v4.md"
DOCX_PATH = HERE / "cadrage_v4.docx"


def _add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(text, style=style)
    return p


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = table.cell(i, j)
            tc.text = cell
            # Header row in bold
            if i == 0:
                for p in tc.paragraphs:
                    for run in p.runs:
                        run.bold = True


def _parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Lit un tableau markdown à partir de la ligne `start` (header).
    Renvoie (rows, next_index)."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # Skip separator row (|---|---|...)
        if re.fullmatch(r"\|[\s\-:|]+\|?", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build_docx() -> None:
    doc = Document()
    # Style base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)

    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        # Tables
        elif line.startswith("|") and "|" in line[1:]:
            rows, next_i = _parse_markdown_table(lines, i)
            _add_table(doc, rows)
            i = next_i
            continue
        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph()
        # Blank
        elif not line.strip():
            pass
        # Code block boundary
        elif line.startswith("```"):
            # Skip until closing ```
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            for bl in block:
                p = doc.add_paragraph(bl, style="Normal")
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
        # Bullets
        elif re.match(r"^\s*[-*]\s+", line):
            txt = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(txt, style="List Bullet")
        else:
            # Inline markdown : on enlève juste les marqueurs simples
            txt = line
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)  # bold
            txt = re.sub(r"`([^`]+)`", r"\1", txt)        # inline code
            if txt.strip():
                doc.add_paragraph(txt)

        i += 1

    doc.save(DOCX_PATH)
    print(f"Cadrage v4 DOCX généré : {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
