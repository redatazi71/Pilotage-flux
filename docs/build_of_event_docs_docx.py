"""Génère les versions DOCX téléchargeables des 3 documents cadrage
OF+EVENT :

  - docs/cadrage_of_event.md → docs/cadrage_of_event.docx
  - docs/cahier_des_charges_of_event.md → docs/cahier_des_charges_of_event.docx
  - docs/user_stories_of_event.md → docs/user_stories_of_event.docx

Réutilise le convertisseur Markdown → DOCX du builder cadrage v4.

Usage : python docs/build_of_event_docs_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


HERE = Path(__file__).resolve().parent

DOCS = [
    ("cadrage_of_event.md", "cadrage_of_event.docx"),
    ("cahier_des_charges_of_event.md", "cahier_des_charges_of_event.docx"),
    ("user_stories_of_event.md", "user_stories_of_event.docx"),
]


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = table.cell(i, j)
            tc.text = cell
            if i == 0:
                for p in tc.paragraphs:
                    for run in p.runs:
                        run.bold = True


def _parse_markdown_table(
    lines: list[str], start: int,
) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.fullmatch(r"\|[\s\-:|]+\|?", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md = md_path.read_text(encoding="utf-8")
    lines = md.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("|") and "|" in line[1:]:
            rows, next_i = _parse_markdown_table(lines, i)
            _add_table(doc, rows)
            i = next_i
            continue
        elif line.strip() == "---":
            doc.add_paragraph()
        elif not line.strip():
            pass
        elif line.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            for bl in block:
                p = doc.add_paragraph(bl, style="Normal")
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
        elif re.match(r"^\s*[-*]\s+", line):
            txt = re.sub(r"^\s*[-*]\s+", "", line)
            # Preserve bold markers as bold runs
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, txt)
        else:
            txt = line
            if txt.strip():
                p = doc.add_paragraph()
                _add_inline(p, txt)

        i += 1

    doc.save(docx_path)
    print(f"[OK] {docx_path.name}")


def _add_inline(paragraph, text: str) -> None:
    """Ajoute du texte avec inline bold (**...**) et inline code (`...`)."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def main() -> int:
    for md_name, docx_name in DOCS:
        md_path = HERE / md_name
        docx_path = HERE / docx_name
        if not md_path.exists():
            print(f"[SKIP] {md_name} introuvable")
            continue
        build_docx(md_path, docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
