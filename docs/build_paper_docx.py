"""Convertit paper_rfgi_v1.md en DOCX téléchargeable, avec :
  - titres hiérarchiques Word (Heading 1..3)
  - tables markdown → tables Word
  - images ![](path) → images embarquées
  - listes à puces
  - conservation du texte gras/italique inline
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt


SOURCE = Path("docs/paper_rfgi_v1.md")
OUT = Path("docs/paper_rfgi_v1.docx")


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")


def _add_inline_runs(paragraph, text: str) -> None:
    """Applique bold / italic / code sur une seule ligne."""
    pos = 0
    tokens: list[tuple[str, str]] = []

    def _find(pattern, style):
        return [(m.start(), m.end(), m.group(1), style)
                for m in pattern.finditer(text)]

    matches = (
        _find(INLINE_BOLD, "bold")
        + _find(INLINE_ITALIC, "italic")
        + _find(INLINE_CODE, "code")
    )
    matches.sort(key=lambda x: x[0])

    for start, end, inner, style in matches:
        if start > pos:
            paragraph.add_run(text[pos:start])
        run = paragraph.add_run(inner)
        if style == "bold":
            run.bold = True
        elif style == "italic":
            run.italic = True
        elif style == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0]
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(h.strip())
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c >= len(header):
                continue
            cell = table.rows[r].cells[c]
            _add_inline_runs(cell.paragraphs[0], val.strip())


def _add_image(doc: Document, path: Path) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Image manquante : {path}]").italic = True
        return
    doc.add_picture(str(path), width=Inches(5.8))


def build_docx() -> None:
    md_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    in_list = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table : détecter en-tête + séparateur | --- |
        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip())
        ):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            _add_table(doc, header, rows)
            doc.add_paragraph("")
            in_list = False
            continue

        # Image ![...](path)
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            caption, ipath = m.group(1), m.group(2)
            _add_image(doc, Path("docs") / ipath)
            if caption:
                cap = doc.add_paragraph(caption)
                cap.style = doc.styles["Caption"]
            i += 1
            in_list = False
            continue

        # Titre
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).rstrip("#").strip()
            heading = doc.add_heading(text, level=min(level, 4))
            i += 1
            in_list = False
            continue

        # Ligne horizontale
        if re.match(r"^-{3,}$", stripped):
            doc.add_paragraph().add_run().add_break()
            i += 1
            in_list = False
            continue

        # Liste
        m = re.match(r"^-\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1))
            in_list = True
            i += 1
            continue

        # Bloc de code triple-backtick — collecter et ignorer les mermaid
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # consommer le closing ```
            if block:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(block))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            in_list = False
            continue

        # Paragraphe normal
        if stripped == "":
            in_list = False
            i += 1
            continue
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"[ok] {OUT}")


if __name__ == "__main__":
    build_docx()
