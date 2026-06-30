"""Génère docs/paper_hal_v1.docx à partir de docs/paper_hal_v1.md.

Réutilise les helpers de `build_cadrage_v4_docx.py` mais cible un
fichier source différent. Format DOCX HAL préformaté : titre, abstract,
sections, biblio.

Usage : python docs/build_paper_hal_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "paper_hal_v1.md"
DOCX_PATH = HERE / "paper_hal_v1.docx"


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = table.cell(i, j)
            tc.text = cell
            if i == 0:
                for p in tc.paragraphs:
                    for run in p.runs:
                        run.bold = True


def _parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.fullmatch(r"\|[\s\-:|]+\|?", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build_docx() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)

    md = MD_PATH.read_text(encoding="utf-8")
    lines = md.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.+?)\)", line)
            if m:
                alt, rel = m.group(1), m.group(2)
                img_path = (HERE / rel).resolve()
                if img_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(6.0))
                    if alt:
                        cap = doc.add_paragraph(alt)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cap.runs:
                            r.italic = True
                            r.font.size = Pt(9)
        elif line.startswith("|") and "|" in line[1:]:
            rows, next_i = _parse_markdown_table(lines, i)
            _add_table(doc, rows)
            i = next_i
            continue
        elif line.strip() == "---":
            doc.add_paragraph()
        elif not line.strip():
            pass
        elif line.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            for bl in block:
                p = doc.add_paragraph(bl, style="Normal")
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
        elif re.match(r"^\s*[-*]\s+", line):
            txt = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(txt, style="List Bullet")
        else:
            txt = line
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)
            txt = re.sub(r"`([^`]+)`", r"\1", txt)
            txt = re.sub(r"\*(.+?)\*", r"\1", txt)  # italic
            if txt.strip():
                doc.add_paragraph(txt)

        i += 1

    doc.save(DOCX_PATH)
    print(f"Paper HAL DOCX généré : {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
