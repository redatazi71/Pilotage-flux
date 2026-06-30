"""Rapport étude paires de domaines — quel duo de domaines stresse le
plus la chaîne BCE en rupture et récupération.

Compile :
  - Matrice 5×5 des paires avec OTIF moyen et recovery_days moyen
  - Heatmap rupture (1 - OTIF) sur la grille 5×5
  - Heatmap récupération sur la grille 5×5
  - Top 10 paires les plus destructrices
  - Lecture comparative (paires homogènes vs hétérogènes)

Usage :
    python docs/build_rapport_pair_docx.py
"""

from __future__ import annotations

import csv
import json
import statistics
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt    # noqa: E402
import numpy as np                   # noqa: E402
from docx import Document            # noqa: E402
from docx.shared import Inches, Pt   # noqa: E402

from pilotage_flux.comparative.domain_pair_stress import DOMAINS

CSV_PATH = Path("docs/domain_pair_runs.csv")
RANK_PATH = Path("docs/domain_pair_ranking.json")
FIGURES_DIR = Path("docs/figures_pair")
FIGURES_DIR.mkdir(exist_ok=True)


def load_runs() -> list[dict]:
    if not CSV_PATH.exists():
        return []
    with CSV_PATH.open() as f:
        return list(csv.DictReader(f))


def load_ranking() -> list[dict]:
    if not RANK_PATH.exists():
        return []
    return json.loads(RANK_PATH.read_text())


def _build_grid(
    ranking: list[dict], key: str,
) -> np.ndarray:
    grid = np.full((len(DOMAINS), len(DOMAINS)), np.nan)
    for i, d_a in enumerate(DOMAINS):
        for j, d_b in enumerate(DOMAINS):
            for e in ranking:
                if e["domain_a"] == d_a and e["domain_b"] == d_b:
                    grid[i][j] = float(e[key])
                    break
    return grid


def fig_heatmap(
    ranking: list[dict], key: str, label: str,
    *, reverse: bool = False, fmt: str = ".2f",
) -> Path:
    """Heatmap 5×5 d'un KPI sur les paires."""
    grid = _build_grid(ranking, key)
    fig, ax = plt.subplots(figsize=(7, 5.5))
    cmap = "RdYlGn_r" if reverse else "RdYlGn"
    im = ax.imshow(grid, cmap=cmap, aspect="auto")
    ax.set_xticks(range(len(DOMAINS)))
    ax.set_yticks(range(len(DOMAINS)))
    ax.set_xticklabels(DOMAINS, rotation=30, ha="right")
    ax.set_yticklabels(DOMAINS)
    ax.set_xlabel("Domaine B")
    ax.set_ylabel("Domaine A")
    ax.set_title(f"{label} — matrice 5×5 paires de domaines")
    # Annote chaque cellule
    for i in range(len(DOMAINS)):
        for j in range(len(DOMAINS)):
            if not np.isnan(grid[i][j]):
                ax.text(
                    j, i, f"{grid[i][j]:{fmt}}",
                    ha="center", va="center",
                    fontsize=9,
                    color="white" if cmap == "RdYlGn_r" and grid[i][j] > np.nanmedian(grid)
                          else "black",
                )
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    path = FIGURES_DIR / f"heatmap_{key}.png"
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _p(doc, text):
    doc.add_paragraph(text)


def _add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)


def build_docx(out_path: Path) -> None:
    runs = load_runs()
    ranking = load_ranking()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    today = datetime.now().strftime("%Y-%m-%d")
    title = doc.add_heading(
        "Étude paires de domaines — rupture × récupération", level=0,
    )
    title.alignment = 1
    p = doc.add_paragraph(
        f"Matrice 5×5 des paires (D_a, D_b) sur la chaîne BCE. "
        f"Daté du {today}."
    )
    p.alignment = 1

    if not ranking:
        _h(doc, "⚠ Étude non exécutée", 1)
        _p(doc,
           "Lancer `python docs/run_domain_pair_study.py --seeds 10` "
           "puis régénérer ce rapport.")
        doc.save(out_path)
        print(f"DOCX écrit (vide) : {out_path}")
        return

    n_ok = sum(1 for r in runs if r.get("status") == "ok")
    n_crashed = sum(1 for r in runs if r.get("status") == "crashed")

    _h(doc, "1. Protocole", 1)
    _p(doc,
       f"Question scientifique : parmi les 25 paires possibles "
       f"(5 domaines × 5 domaines), quelle paire de domaines a le "
       f"plus grand impact en rupture (OTIF) et récupération "
       f"(temps de retour WIP) ?")
    _p(doc,
       "Méthodologie : pour chaque paire (D_a, D_b), 12 hazards "
       "concentrés (6 du domaine D_a, 6 du domaine D_b — 12 du même "
       "si D_a = D_b) sont injectés sur un horizon de 60 jours avec "
       "espacement minimal de 3 jours et payload aléatoire selon "
       "seed. Pilotage de référence : FLUX+EVENT+BCE (le plus complet "
       "de la couche cybernétique). Saturation 0.94 (zone de stress "
       "proche rupture). N seeds = 5.")
    _p(doc,
       f"Volumétrie : {len(runs)} runs, {n_ok} ok, {n_crashed} crashs.")

    _h(doc, "2. Heatmaps des KPIs sur la grille 5×5", 1)
    _p(doc, "Vert = bon, rouge = mauvais. La diagonale représente "
            "les stress mono-domaine (12 hazards d'un seul kind).")

    fig_otif = fig_heatmap(
        ranking, "otif_mean", "OTIF moyen",
        reverse=False, fmt=".1%",
    )
    _p(doc, "Figure 1 — OTIF moyen (haut = meilleur).")
    doc.add_picture(str(fig_otif), width=Inches(6.0))

    fig_recov = fig_heatmap(
        ranking, "recovery_mean_days", "Temps de récupération (j)",
        reverse=True, fmt=".1f",
    )
    _p(doc, "Figure 2 — Temps moyen de retour du WIP dans la bande "
            "post-hazard (bas = meilleur).")
    doc.add_picture(str(fig_recov), width=Inches(6.0))

    fig_score = fig_heatmap(
        ranking, "combined_impact_score", "Score combiné destructivité",
        reverse=True, fmt=".3f",
    )
    _p(doc, "Figure 3 — Score combiné = (1 − OTIF) + recovery/30. "
            "Plus le score est élevé, plus la paire est destructrice.")
    doc.add_picture(str(fig_score), width=Inches(6.0))

    _h(doc, "3. Ranking des 10 paires les plus destructrices", 1)
    rows = [["Rang", "Paire (D_a | D_b)", "OTIF",
              "Recovery (j)", "Yield", "Score"]]
    for i, e in enumerate(ranking[:10], start=1):
        rows.append([
            str(i),
            e["pair"],
            f"{float(e['otif_mean']):.1%}",
            f"{float(e['recovery_mean_days']):.2f}",
            f"{float(e['yield_mean']):.1%}",
            f"{float(e['combined_impact_score']):.4f}",
        ])
    _add_table(doc, rows)

    _h(doc, "4. Lecture comparative — homogène vs hétérogène", 1)
    # Sépare diagonale et hors-diagonale
    diag = [e for e in ranking if e["domain_a"] == e["domain_b"]]
    off_diag = [e for e in ranking if e["domain_a"] != e["domain_b"]]
    if diag and off_diag:
        diag_otif = statistics.mean(
            float(e["otif_mean"]) for e in diag
        )
        off_otif = statistics.mean(
            float(e["otif_mean"]) for e in off_diag
        )
        diag_rec = statistics.mean(
            float(e["recovery_mean_days"]) for e in diag
        )
        off_rec = statistics.mean(
            float(e["recovery_mean_days"]) for e in off_diag
        )
        _p(doc,
           f"OTIF moyen paires homogènes (D,D) : {diag_otif:.1%}. "
           f"OTIF moyen paires hétérogènes : {off_otif:.1%}. "
           f"Différence : {(off_otif - diag_otif):+.1%}.")
        _p(doc,
           f"Récupération paires homogènes : {diag_rec:.2f} j. "
           f"Hétérogènes : {off_rec:.2f} j. "
           f"Différence : {(off_rec - diag_rec):+.2f} j.")
        if off_otif < diag_otif:
            _p(doc,
               "→ Les paires **hétérogènes** sont plus destructrices "
               "que les paires mono-domaine. La doctrine cybernétique "
               "absorbe mieux un stress homogène (qui touche peu de "
               "cellules MACRS) qu'un stress hétérogène (qui touche "
               "des cellules de domaines disjoints, plus difficiles "
               "à corréler).")
        else:
            _p(doc,
               "→ Les paires **homogènes** sont plus destructrices. "
               "La concentration de hazards sur un seul domaine "
               "épuise les degrés de liberté d'une racine MACRS et "
               "fait converger le système vers la rupture.")

    _h(doc, "5. Top 3 et lecture cybernétique", 1)
    _p(doc,
       "Pour les 3 paires les plus destructrices, on rappelle "
       "les racines MACRS doctrinalement touchées :")
    racine_by_domain = {
        "demande":          "R005 Avance de commande",
        "approvisionnement": "R011 Retard livraison fournisseur",
        "logistique":       "R019 Incident transport interne",
        "production":       "R030 Panne machine",
        "qualite":          "R039 NC produit interne",
    }
    rows = [["Rang", "Paire", "Racine A", "Racine B"]]
    for i, e in enumerate(ranking[:3], start=1):
        rows.append([
            str(i),
            e["pair"],
            racine_by_domain.get(e["domain_a"], "?"),
            racine_by_domain.get(e["domain_b"], "?"),
        ])
    _add_table(doc, rows)

    _h(doc, "6. Limites de l'étude", 1)
    _p(doc,
       "(a) Saturation fixe 0.94 : à d'autres niveaux le ranking "
       "peut bouger. (b) Pilotage de référence unique "
       "(FLUX+EVENT+BCE) ; refaire l'étude sur OF montrerait quel "
       "stress résiste à la couche cybernétique. (c) N seeds = 5 ; "
       "passer à 20+ resserrerait les barres d'erreur. (d) Les "
       "domaines sont attaqués par UN seul kind chacun ; un domaine "
       "comme « production » a en réalité 12 racines distinctes "
       "(R027..R038) et notre étude n'en stresse qu'une seule.")

    doc.save(out_path)
    print(f"DOCX écrit : {out_path}")


def main() -> int:
    out = Path("docs/rapport_pairs.docx")
    build_docx(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
