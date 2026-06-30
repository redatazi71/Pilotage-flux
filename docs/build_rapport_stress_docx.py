"""Rapport stress — section "Résultats scientifiques" en DOCX.

Compile :
  - 5 questions scientifiques avec verdicts
  - Tableaux de stats par KPI
  - Dashboard 6 pilotages × 6 KPIs (figures)
  - Robustesse par pilotage

Usage :
    python docs/build_rapport_stress_docx.py
"""

from __future__ import annotations

import csv
import json
import statistics
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt    # noqa: E402
from docx import Document            # noqa: E402
from docx.shared import Inches, Pt   # noqa: E402


CSV_PATH = Path("docs/stress_study_runs.csv")
STATS_PATH = Path("docs/stress_study_stats.json")
FIGURES_DIR = Path("docs/figures_stress")
FIGURES_DIR.mkdir(exist_ok=True)


def load_runs() -> list[dict]:
    if not CSV_PATH.exists():
        return []
    with CSV_PATH.open() as f:
        return list(csv.DictReader(f))


def load_stats() -> list[dict]:
    if not STATS_PATH.exists():
        return []
    return json.loads(STATS_PATH.read_text())


# ---------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------


def fig_kpi_by_pilotage_saturation(
    runs: list[dict], kpi: str, ylabel: str,
    *, higher_is_better: bool = True,
) -> Path:
    by_pil_sat: dict[tuple, list[float]] = defaultdict(list)
    for r in runs:
        if r.get("status") != "ok":
            continue
        try:
            sat = float(r["saturation"])
            val = float(r[kpi])
        except (KeyError, ValueError):
            continue
        by_pil_sat[(r["doctrine"], sat)].append(val)

    pilotages = sorted({k[0] for k in by_pil_sat})
    saturations = sorted({k[1] for k in by_pil_sat})

    fig, ax = plt.subplots(figsize=(9, 5))
    colors = {
        "of": "#888", "of_event": "#1f77b4",
        "of_event_bce": "#2ca02c",
        "flux": "#ff7f0e", "event": "#d62728",
        "event_bce": "#9467bd",
    }
    for pil in pilotages:
        xs, ys, es = [], [], []
        for s in saturations:
            vals = by_pil_sat.get((pil, s))
            if not vals:
                continue
            xs.append(s)
            ys.append(statistics.mean(vals))
            es.append(
                statistics.stdev(vals) if len(vals) >= 2 else 0.0
            )
        marker = "o" if "bce" in pil else "s"
        ax.errorbar(
            xs, ys, yerr=es, marker=marker, label=pil, linewidth=2,
            color=colors.get(pil), capsize=3, alpha=0.85,
        )
    ax.set_xlabel("Saturation R1")
    ax.set_ylabel(ylabel + (" (haut = meilleur)" if higher_is_better
                              else " (bas = meilleur)"))
    ax.set_title(f"{ylabel} par pilotage et saturation (stress)")
    ax.axvline(1.0, color="red", linestyle="--", alpha=0.5)
    ax.legend(loc="best", fontsize=8)
    ax.grid(True, alpha=0.3)
    path = FIGURES_DIR / f"{kpi}_by_pilotage_saturation.png"
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def fig_n1_n4_distribution(runs: list[dict]) -> Path:
    """Distribution N1..N4 pour les pilotages BCE — vérifie que la
    recalibration produit la doctrine attendue (N1/N2 majoritaire)."""
    by_pil = defaultdict(lambda: defaultdict(list))
    for r in runs:
        if r.get("status") != "ok":
            continue
        if not r["doctrine"].endswith("_bce"):
            continue
        for n in ("n1", "n2", "n3", "n4"):
            try:
                by_pil[r["doctrine"]][n].append(int(r[n]))
            except (KeyError, ValueError):
                pass

    pilotages = sorted(by_pil)
    if not pilotages:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.text(0.5, 0.5, "Pas de données BCE", ha="center", va="center",
                 transform=ax.transAxes)
        path = FIGURES_DIR / "n1_n4_distribution.png"
        fig.savefig(path, dpi=150)
        plt.close(fig)
        return path

    means = {n: [] for n in ("n1", "n2", "n3", "n4")}
    for pil in pilotages:
        for n in ("n1", "n2", "n3", "n4"):
            vals = by_pil[pil][n]
            means[n].append(statistics.mean(vals) if vals else 0)

    fig, ax = plt.subplots(figsize=(7, 4.5))
    bottom = [0.0] * len(pilotages)
    colors = {"n1": "#9ecae1", "n2": "#6baed6",
              "n3": "#fd8d3c", "n4": "#a63603"}
    labels = {"n1": "N1 absorption", "n2": "N2 ajust auto",
              "n3": "N3 replan local", "n4": "N4 replan global"}
    for n in ("n1", "n2", "n3", "n4"):
        ax.bar(pilotages, means[n], bottom=bottom,
                label=labels[n], color=colors[n])
        bottom = [b + v for b, v in zip(bottom, means[n])]
    ax.set_ylabel("Nb moyen de delta_decisions par run")
    ax.set_title("Distribution N1..N4 — doctrine cadrage v1.3 §3.11")
    ax.legend()
    ax.grid(True, alpha=0.3, axis="y")
    path = FIGURES_DIR / "n1_n4_distribution.png"
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


# ---------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _p(doc, text):
    doc.add_paragraph(text)


def _add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)


def build_docx(out_path: Path) -> None:
    runs = load_runs()
    stats = load_stats()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    today = datetime.now().strftime("%Y-%m-%d")
    title = doc.add_heading(
        "Résultats scientifiques — étude stress BCE", level=0,
    )
    title.alignment = 1
    p = doc.add_paragraph(
        f"Comparaison de 6 pilotages APS+MES en conditions stress. "
        f"Daté du {today}."
    )
    p.alignment = 1

    if not runs:
        _h(doc, "⚠ Étude non exécutée", 1)
        _p(doc,
           "Lancer `python docs/run_stress_study.py --seeds 10` "
           "puis régénérer ce rapport.")
        doc.save(out_path)
        print(f"DOCX écrit (vide) : {out_path}")
        return

    n_ok = sum(1 for r in runs if r.get("status") == "ok")
    n_crashed = sum(1 for r in runs if r.get("status") == "crashed")
    _h(doc, "1. Protocole expérimental", 1)
    _p(doc,
       f"Plan d'expérience : 6 pilotages × 6 saturations × 3 "
       f"implantations × N seeds. Scénario stress 60 jours, 12 "
       f"hazards diversifiés (4 BREAKDOWN, 3 QUALITY_NC, 2 PO_DELAY, "
       f"2 URGENT_ORDER, 1 LOGISTIC_DELAY). Saturations couvrent "
       f"78% à 110% (intentionnellement au-delà de 100% pour forcer "
       f"la rupture). Jitter ±2 jours + ±20% sur payload selon seed.")
    _p(doc, f"Volumétrie : {len(runs)} runs total, {n_ok} ok, "
            f"{n_crashed} crashs.")
    _p(doc,
       "Profils filtre dual : DEFAULT historique (seuils "
       "0.20/0.50/1.00/2.00/3.50) pour les pilotages OF+EVENT et "
       "FLUX+EVENT ; CONSERVATIVE (0.50/1.00/1.50/2.00/3.00) pour "
       "les pilotages BCE — calibration doctrinale qui favorise "
       "l'absorption N1/N2 plutôt que le replan N3/N4.")

    _h(doc, "2. Dashboard 6 KPIs × pilotages × saturations", 1)
    for kpi, ylab, hib in (
        ("otif", "OTIF", True),
        ("yield_pct", "Yield", True),
        ("wip_mean", "WIP moyen", False),
        ("lateness_mean_days", "Retard moyen (j)", False),
        ("cost_per_good_unit", "Coût/unité bonne (€)", False),
        ("mean_recovery_days", "Temps de récupération (j)", False),
    ):
        path = fig_kpi_by_pilotage_saturation(
            runs, kpi, ylab, higher_is_better=hib,
        )
        _p(doc, f"Figure : {ylab} par pilotage et saturation.")
        doc.add_picture(str(path), width=Inches(6.0))

    _h(doc, "3. Distribution N1..N4 doctrinale", 1)
    _p(doc,
       "Cadrage v1.3 §3.11 prévoit une distribution avec N1/N2 "
       "majoritaires (absorption + ajustement auto) et N3/N4 "
       "exceptionnels (replan humain). La calibration CONSERVATIVE "
       "du profil filtre dual sur les pilotages BCE doit produire "
       "cette distribution.")
    path = fig_n1_n4_distribution(runs)
    doc.add_picture(str(path), width=Inches(6.0))

    _h(doc, "4. Résultats des 5 questions scientifiques", 1)

    questions_by_label: dict[str, list[dict]] = defaultdict(list)
    rob_stat = None
    for s in stats:
        if "by_pilotage" in s:
            rob_stat = s
            continue
        # Q1..Q4 ont un label "Qx (kpi)"
        qid = s["label"].split(" ")[0]
        questions_by_label[qid].append(s)

    q_titles = {
        "Q1": "Q1. Apport EVENT — V0 OF + event sourcing améliore-t-il les KPIs ?",
        "Q2": "Q2. Apport FLUX — V2 contractualisation flux améliore-t-elle vs OF+EVENT ?",
        "Q3": "Q3. Apport BCE sur FLUX — la couche cybernétique améliore-t-elle FLUX+EVENT ?",
        "Q4": "Q4. Apport BCE sur OF — la couche cybernétique améliore-t-elle OF+EVENT ?",
    }

    for qid in ("Q1", "Q2", "Q3", "Q4"):
        if qid not in questions_by_label:
            continue
        _h(doc, q_titles[qid], 2)
        rows = [["KPI", "n", "Δ médiane (a−b)",
                  "CI 95% bas", "CI 95% haut",
                  "Cliff's δ", "Wilcoxon p", "Verdict"]]
        for s in questions_by_label[qid]:
            kpi_short = s["label"].split("(", 1)[1].rstrip(")")
            rows.append([
                kpi_short,
                str(s.get("n_pairs", 0)),
                f"{s.get('median_diff', 0):+.4f}",
                f"{s.get('ci_low', 0):+.4f}",
                f"{s.get('ci_high', 0):+.4f}",
                f"{s.get('cliffs_delta', 0):+.3f}",
                (f"{s.get('wilcoxon_p'):.4f}"
                 if s.get("wilcoxon_p") is not None else "n/a"),
                s.get("verdict", "n/a"),
            ])
        _add_table(doc, rows)

    if rob_stat:
        _h(doc, "Q5. Robustesse — seuil de rupture OTIF par pilotage", 2)
        _p(doc,
           "La saturation à laquelle l'OTIF moyen franchit le seuil "
           "0.90 à la baisse (interpolation linéaire). Valeur élevée "
           "= système robuste sous charge.")
        rob = rob_stat["by_pilotage"]
        rows = [["Pilotage", "Saturation de rupture"]]
        for pil in sorted(rob):
            bp = rob[pil]
            v = (f"{float(bp):.2f}" if bp is not None
                  else "robuste partout")
            rows.append([pil, v])
        _add_table(doc, rows)

    _h(doc, "5. Lecture comparative", 1)
    _p(doc,
       "Les comparaisons sont **appariées** par (saturation, "
       "implantation, seed) ce qui contrôle pour l'effet de "
       "scénario : on compare 2 pilotages sur exactement le même "
       "contexte. Le **Wilcoxon signed-rank** teste si la médiane "
       "des différences est significativement différente de 0 "
       "(p-value bilatérale via approximation normale). Le "
       "**bootstrap percentile** (1000 rééchantillonnages) donne "
       "l'intervalle de confiance à 95% sur la médiane des "
       "différences. Le **Cliff's δ** mesure l'effet sans hypothèse "
       "normale : valeur dans [-1, 1], positive si le premier "
       "pilotage produit des valeurs plus élevées en moyenne.")
    _p(doc,
       "Pour les KPIs où plus est mieux (OTIF, yield, robustesse), "
       "un Δ médian positif et Cliff's δ > 0 indiquent un gain. "
       "Pour les KPIs où moins est mieux (WIP, retard, coût, temps "
       "de récupération), c'est l'inverse.")

    _h(doc, "6. Limites du présent résultat", 1)
    _p(doc,
       "(a) Scénario stress reste mono-article (ART-A) sur "
       "fixtures_extended — un mix multi-articles avec saisonnalité "
       "produirait plus de variance et discriminerait davantage les "
       "pilotages flux vs OF. (b) Saturation 110% calibrée par "
       "volume SO en linéaire ; en pratique d'autres mécaniques "
       "(absence opérateur, casse outillage) pourraient amplifier "
       "ou atténuer l'effet. (c) Profil CONSERVATIVE BCE choisi "
       "doctrinalement mais d'autres calibrations sont possibles ; "
       "une étude d'hyperparamètres sur les seuils donnerait la "
       "frontière de Pareto vraie. (d) N seeds modeste (≥ 5) — "
       "augmenter à 20+ resserrerait les CI.")

    doc.save(out_path)
    print(f"DOCX écrit : {out_path}")


def main() -> int:
    out = Path("docs/rapport_stress.docx")
    build_docx(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
